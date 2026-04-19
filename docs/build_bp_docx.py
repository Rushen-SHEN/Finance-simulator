"""
Pipeline: ARIA_BP_External.md → clean MD → pandoc → styled DOCX.
Adapted from skills/docx-pipeline for the ARIA BP External document.
Content is NEVER modified — only Markdown escape artifacts are stripped,
then proper conversion and professional styling applied.
"""

import re, subprocess, sys, os

# ── Paths ───────────────────────────────────────────────────────
BASE = os.path.dirname(os.path.abspath(__file__))
SRC_MD   = os.path.join(BASE, "ARIA_BP_External.md")
CLEAN_MD = os.path.join(BASE, "ARIA_BP_External_clean.md")
RAW_DOCX = os.path.join(BASE, "ARIA_BP_External_raw.docx")
OUT_DOCX = os.path.join(BASE, "ARIA_BP_External.docx")


def is_table_sep_row(line):
    """Check if a line is a Markdown table separator row like |---|:---:|."""
    s = line.strip()
    if not (s.startswith('|') and s.endswith('|')):
        return False
    cells = s[1:-1].split('|')
    for cell in cells:
        c = cell.strip()
        if not c:
            return False
        if not re.match(r'^:?-{1,}:?$', c):
            return False
    return True


def convert_simple_tables_to_pipe(text):
    """Convert pandoc simple/grid tables (space-aligned, dash borders) to pipe tables."""
    lines = text.split('\n')
    result = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if re.match(r'^\s*-{20,}\s*$', stripped):
            table_lines = [line]
            j = i + 1
            found_bottom = False
            while j < len(lines):
                table_lines.append(lines[j])
                tl_stripped = lines[j].strip()
                if re.match(r'^\s*-{20,}\s*$', tl_stripped) and j > i + 2:
                    found_bottom = True
                    break
                j += 1

            if not found_bottom:
                result.append(line)
                i += 1
                continue

            sep_line_idx = None
            for k in range(1, len(table_lines) - 1):
                tl = table_lines[k].rstrip()
                if re.match(r'^\s+(-{2,}\s+){2,}-{2,}\s*$', tl):
                    sep_line_idx = k
                    break

            if sep_line_idx is None:
                result.append(line)
                i += 1
                continue

            sep_text = table_lines[sep_line_idx]
            col_spans = []
            in_dash = False
            start = 0
            for ci, ch in enumerate(sep_text):
                if ch == '-':
                    if not in_dash:
                        start = ci
                        in_dash = True
                else:
                    if in_dash:
                        col_spans.append((start, ci))
                        in_dash = False
            if in_dash:
                col_spans.append((start, len(sep_text)))

            if len(col_spans) < 2:
                result.append(line)
                i += 1
                continue

            def extract_cells_smart(text_line):
                stripped_line = text_line.strip()
                parts = re.split(r'\s{2,}', stripped_line)
                return [p.strip() for p in parts if p.strip()]

            header_cells = []
            for k in range(1, sep_line_idx):
                row_text = table_lines[k]
                if row_text.strip():
                    header_cells = extract_cells_smart(row_text)
                    break

            data_rows = []
            for k in range(sep_line_idx + 1, len(table_lines) - 1):
                row_text = table_lines[k]
                if row_text.strip():
                    data_rows.append(extract_cells_smart(row_text))

            if not header_cells:
                result.append(line)
                i += 1
                continue

            ncols = len(header_cells)
            pipe_lines = []
            pipe_lines.append('| ' + ' | '.join(header_cells) + ' |')
            pipe_lines.append('| ' + ' | '.join(['---'] * ncols) + ' |')
            for row in data_rows:
                while len(row) < ncols:
                    row.append('')
                if len(row) > ncols:
                    row = row[:ncols-1] + [' '.join(row[ncols-1:])]
                pipe_lines.append('| ' + ' | '.join(row) + ' |')

            for pl in pipe_lines:
                result.append(pl)
            i = j + 1
            continue

        result.append(line)
        i += 1

    return '\n'.join(result)


def preprocess_md(src, dst):
    """Remove backslash escapes that prevent pandoc from parsing MD syntax."""
    with open(src, "r", encoding="utf-8") as f:
        text = f.read()

    text = convert_simple_tables_to_pipe(text)

    # Fix doubled heading markers
    text = re.sub(r'^# \\## ', '# ', text, flags=re.MULTILINE)
    text = re.sub(r'^## \\### ', '## ', text, flags=re.MULTILINE)
    text = re.sub(r'^\\### ', '### ', text, flags=re.MULTILINE)
    text = re.sub(r'^\\## ', '## ', text, flags=re.MULTILINE)

    # Fix horizontal rules
    text = re.sub(r'^\\-\\-\\-', '---', text, flags=re.MULTILINE)
    text = text.replace('\\-\\--', '---')

    # Fix escaped pipes
    text = text.replace('\\|', '|')

    # Fix bold markers
    text = text.replace('\\*\\*', '**')
    text = text.replace('\\*', '*')

    # Fix escaped quotes
    text = text.replace('\\\\\\"', '"')
    text = text.replace('\\"', '"')

    # Fix escaped brackets
    text = text.replace('\\[', '[')
    text = text.replace('\\]', ']')

    # Fix math delimiters
    text = text.replace('#\\$#', '$')
    text = text.replace('#$#', '$')

    # Fix LaTeX escapes
    text = text.replace('\\\\frac', '\\frac')
    text = text.replace('\\\\text', '\\text')

    # Fix escaped dollar sign
    text = text.replace('\\$', '$')

    # Fix em-dashes
    text = text.replace('------', '——')

    # Convert en-dashes in non-table lines
    lines = text.split('\n')
    processed = []
    for line in lines:
        stripped = line.strip()
        if is_table_sep_row(stripped):
            processed.append(line)
        elif stripped == '---':
            processed.append(line)
        else:
            processed.append(re.sub(r'(?<!-)--(?!-)', '–', line))
    lines = processed

    # Normalize table separator rows
    normalized = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = stripped[1:-1].split('|')
            looks_like_sep = all(
                re.match(r'^\s*:?[–—-]+:?\s*$', c) for c in cells if c.strip()
            )
            if looks_like_sep and not is_table_sep_row(stripped):
                fixed_cells = []
                for c in cells:
                    c = c.strip()
                    left_colon = c.startswith(':')
                    right_colon = c.endswith(':')
                    if left_colon and right_colon:
                        fixed_cells.append(':---:')
                    elif right_colon:
                        fixed_cells.append('---:')
                    elif left_colon:
                        fixed_cells.append(':---')
                    else:
                        fixed_cells.append('---')
                normalized.append('| ' + ' | '.join(fixed_cells) + ' |')
                continue
        normalized.append(line)
    lines = normalized

    # Collapse blank lines between table rows
    collapsed = []
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        is_tbl = stripped.startswith('|') and stripped.endswith('|') and stripped.count('|') >= 2

        if is_tbl:
            collapsed.append(lines[i])
            i += 1
            while i < len(lines):
                next_stripped = lines[i].strip()
                if next_stripped == '':
                    j = i + 1
                    while j < len(lines) and lines[j].strip() == '':
                        j += 1
                    if j < len(lines):
                        future = lines[j].strip()
                        if future.startswith('|') and future.endswith('|') and future.count('|') >= 2:
                            i = j
                            continue
                        else:
                            break
                    else:
                        break
                elif next_stripped.startswith('|') and next_stripped.endswith('|'):
                    collapsed.append(lines[i])
                    i += 1
                else:
                    break
        else:
            collapsed.append(lines[i])
            i += 1
    lines = collapsed

    # Ensure blank lines around table blocks
    final_lines = []
    in_table = False
    for i, line in enumerate(lines):
        stripped = line.strip()
        is_tbl = stripped.startswith('|') and stripped.endswith('|') and stripped.count('|') >= 2

        if is_tbl and not in_table:
            if final_lines and final_lines[-1].strip() != '':
                final_lines.append('')
            in_table = True
        elif not is_tbl and in_table:
            if final_lines and final_lines[-1].strip() != '':
                final_lines.append('')
            in_table = False

        final_lines.append(line)

    text = '\n'.join(final_lines)

    with open(dst, "w", encoding="utf-8") as f:
        f.write(text)

    print(f"[1/3] Preprocessed: {src} -> {dst}")
    print(f"      {len(final_lines)} lines processed")


def convert_pandoc(src_md, dst_docx):
    """Run pandoc to convert clean MD -> DOCX."""
    cmd = [
        "pandoc", src_md,
        "-f", "markdown-smart+pipe_tables+grid_tables",
        "-t", "docx",
        "--columns=999",
        "-o", dst_docx
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"[ERROR] pandoc failed: {result.stderr}")
        sys.exit(1)
    size_kb = os.path.getsize(dst_docx) / 1024
    print(f"[2/3] Pandoc converted: {dst_docx} ({size_kb:.1f} KB)")


def style_docx(src_docx, dst_docx):
    """Apply professional investor-grade styling to the DOCX."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    DARK_RED    = RGBColor(0x8B, 0x00, 0x00)
    MEDIUM_RED  = RGBColor(0xA0, 0x20, 0x20)
    DARK_GRAY   = RGBColor(0x33, 0x33, 0x33)
    BODY_COLOR  = RGBColor(0x22, 0x22, 0x22)
    WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
    CN_FONT = "SimSun"
    EN_FONT = "Times New Roman"
    TABLE_HDR_BG = "8B0000"
    TABLE_ALT_BG = "F9F5F5"

    def set_cell_shading(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_run_font(run, cn=CN_FONT, en=EN_FONT, size=None, bold=None, color=None):
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), cn)
        rFonts.set(qn("w:ascii"), en)
        rFonts.set(qn("w:hAnsi"), en)
        rFonts.set(qn("w:cs"), en)
        if size is not None:
            run.font.size = size
        if bold is not None:
            run.font.bold = bold
        if color is not None:
            run.font.color.rgb = color

    def style_paragraph_font(para, size=None, bold=None, color=None):
        for run in para.runs:
            effective_bold = bold if bold is not None else run.font.bold
            set_run_font(run, size=size, bold=effective_bold, color=color)

    def set_paragraph_spacing(para, before=None, after=None, line_spacing=None):
        pf = para.paragraph_format
        if before is not None: pf.space_before = before
        if after is not None: pf.space_after = after
        if line_spacing is not None: pf.line_spacing = line_spacing

    def add_page_break_before(para):
        pPr = para._p.get_or_add_pPr()
        for existing in pPr.findall(qn("w:pageBreakBefore")):
            pPr.remove(existing)
        pPr.append(parse_xml(f'<w:pageBreakBefore {nsdecls("w")} w:val="1"/>'))

    def set_paragraph_border_left(para, color="8B0000", size="18", space="6"):
        pPr = para._p.get_or_add_pPr()
        borders = pPr.find(qn("w:pBdr"))
        if borders is None:
            borders = parse_xml(f'<w:pBdr {nsdecls("w")}/>')
            pPr.append(borders)
        left = parse_xml(
            f'<w:left {nsdecls("w")} w:val="single" w:sz="{size}" '
            f'w:space="{space}" w:color="{color}"/>'
        )
        for existing in borders.findall(qn("w:left")):
            borders.remove(existing)
        borders.append(left)

    def set_style_fonts(style_obj):
        rPr = style_obj.element.find(qn("w:rPr"))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
            style_obj.element.append(rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), cn)
        rFonts.set(qn("w:ascii"), en)
        rFonts.set(qn("w:hAnsi"), en)

    # ── Load ────────────────────────────────────────────────────
    doc = Document(src_docx)

    # ── Page Setup ──────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.8)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    # ── Heading Styles ──────────────────────────────────────────
    styles = doc.styles
    style_names = [s.name for s in styles]

    for name, sz, clr in [("Heading 1", 18, DARK_RED),
                           ("Heading 2", 14, MEDIUM_RED),
                           ("Heading 3", 12, DARK_GRAY)]:
        if name in style_names:
            s = styles[name]
        else:
            s = styles.add_style(name, 1)
        s.font.size = Pt(sz)
        s.font.bold = True
        s.font.color.rgb = clr
        sp_before = {18: 36, 14: 24, 12: 18}[sz]
        sp_after  = {18: 14, 14: 10, 12: 8}[sz]
        s.paragraph_format.space_before = Pt(sp_before)
        s.paragraph_format.space_after = Pt(sp_after)
        s.paragraph_format.line_spacing = 1.3
        set_style_fonts(s)

    # ── Process Paragraphs ──────────────────────────────────────
    first_h1_seen = False
    for para in doc.paragraphs:
        sn = (para.style.name or "").lower()
        level = 0
        if sn.startswith("heading"):
            try: level = int(sn.replace("heading", "").strip())
            except: pass

        if level == 1:
            style_paragraph_font(para, size=Pt(18), bold=True, color=DARK_RED)
            set_paragraph_spacing(para, before=Pt(36), after=Pt(14), line_spacing=1.3)
            if first_h1_seen:
                add_page_break_before(para)
            first_h1_seen = True

        elif level == 2:
            style_paragraph_font(para, size=Pt(14), bold=True, color=MEDIUM_RED)
            set_paragraph_spacing(para, before=Pt(24), after=Pt(10), line_spacing=1.3)

        elif level == 3:
            style_paragraph_font(para, size=Pt(12), bold=True, color=DARK_GRAY)
            set_paragraph_spacing(para, before=Pt(18), after=Pt(8), line_spacing=1.3)
            set_paragraph_border_left(para)

        else:
            text = para.text.strip()
            is_warning = any(kw in text for kw in [
                "⚠", "风险提示", "重要披露", "当前阶段", "事实前提声明",
                "关键假设", "局限性", "本报告所有财务数据"
            ])
            if is_warning:
                style_paragraph_font(para, size=Pt(10), color=DARK_RED)
                set_paragraph_spacing(para, before=Pt(8), after=Pt(8), line_spacing=1.4)
            else:
                style_paragraph_font(para, size=Pt(10.5), color=BODY_COLOR)
                set_paragraph_spacing(para, before=Pt(4), after=Pt(4), line_spacing=1.4)

            if para.style.name and "quote" in para.style.name.lower():
                para.paragraph_format.left_indent = Cm(1.0)
                set_paragraph_border_left(para, color="CCCCCC", size="8", space="8")

    # ── Process Tables ──────────────────────────────────────────
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tblPr = table._tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            table._tbl.insert(0, tblPr)

        borders_xml = (
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'</w:tblBorders>'
        )
        for old in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(old)
        tblPr.append(parse_xml(borders_xml))

        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if row_idx == 0:
                    set_cell_shading(cell, TABLE_HDR_BG)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            set_run_font(run, size=Pt(9.5), bold=True, color=WHITE)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    if row_idx % 2 == 0:
                        set_cell_shading(cell, TABLE_ALT_BG)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            set_run_font(run, size=Pt(9), color=BODY_COLOR)

                tcPr = cell._tc.get_or_add_tcPr()
                mar = parse_xml(
                    f'<w:tcMar {nsdecls("w")}>'
                    f'  <w:top w:w="40" w:type="dxa"/>'
                    f'  <w:left w:w="80" w:type="dxa"/>'
                    f'  <w:bottom w:w="40" w:type="dxa"/>'
                    f'  <w:right w:w="80" w:type="dxa"/>'
                    f'</w:tcMar>'
                )
                for old in tcPr.findall(qn("w:tcMar")):
                    tcPr.remove(old)
                tcPr.append(mar)

                for p in cell.paragraphs:
                    set_paragraph_spacing(p, before=Pt(2), after=Pt(2), line_spacing=1.15)

    # ── Footer ──────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ARIA | Confidential")
        set_run_font(run, size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))

    # ── Save ────────────────────────────────────────────────────
    doc.save(dst_docx)
    size_kb = os.path.getsize(dst_docx) / 1024
    print(f"[3/3] Styled: {dst_docx} ({size_kb:.1f} KB)")


def verify_docx(docx_path):
    """Quick validation — check for MD artifacts, tables, headings."""
    from docx import Document

    doc = Document(docx_path)
    issues = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        if "\\*\\*" in t or "\\|" in t or "\\###" in t or "\\##" in t:
            issues.append(f"  P{i} [escaped]: {t[:90]}")
        stripped = t.strip()
        if stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 3:
            issues.append(f"  P{i} [raw-tbl]: {t[:90]}")
        if "**" in t:
            issues.append(f"  P{i} [literal**]: {t[:90]}")

    if issues:
        print(f"\nVERIFY ISSUES ({len(issues)}):")
        for x in issues[:30]:
            print(x)
    else:
        print("\nVERIFY OK — No MD artifacts in paragraphs")

    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    for i, tbl in enumerate(doc.tables):
        rows = len(tbl.rows)
        cols = len(tbl.columns)
        hdr = " | ".join(c.text[:25] for c in tbl.rows[0].cells)
        print(f"  T{i}: {rows}r x {cols}c — {hdr}")

    h_count = {1: 0, 2: 0, 3: 0}
    for p in doc.paragraphs:
        sn = (p.style.name or "").lower()
        if sn.startswith("heading"):
            try:
                lv = int(sn.replace("heading", "").strip())
                if lv in h_count:
                    h_count[lv] += 1
            except:
                pass
    print(f"\nHeadings: H1={h_count[1]}, H2={h_count[2]}, H3={h_count[3]}")


def main():
    print("=" * 60)
    print("ARIA BP External — MD -> DOCX Professional Pipeline")
    print("=" * 60)
    preprocess_md(SRC_MD, CLEAN_MD)
    convert_pandoc(CLEAN_MD, RAW_DOCX)
    style_docx(RAW_DOCX, OUT_DOCX)
    print("=" * 60)
    print(f"DONE -> {OUT_DOCX}")

    # Verify
    verify_docx(OUT_DOCX)

    # Cleanup intermediate raw docx
    for tmp in [RAW_DOCX]:
        if os.path.exists(tmp):
            os.remove(tmp)
            print(f"\n  Cleaned up: {os.path.basename(tmp)}")
    print(f"  Kept for review: {os.path.basename(CLEAN_MD)}")


if __name__ == "__main__":
    main()
